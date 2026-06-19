#!/usr/bin/env python3
"""
Build the Word deliverables for the second revision from the markdown sources.

Produces, in notebooks/panelfm/revision2/docs/:
  manuscript_clean.docx            final text, no markup (for reading)
  manuscript_bold.docx             final text with new/moved text in **bold** (journal submission)
  manuscript_tracked_changes.docx  genuine Word redline (R1 -> R2) via w:ins / w:del
  supplementary_appendix.docx      clean appendix
  response_to_reviewers.docx       blinded response letter

The bold source (manuscript_bold.md) marks revised passages with **...**. The clean
source is the bold source with markup stripped. The tracked-changes redline is a
paragraph-aligned word-level diff of the R1 clean manuscript against the R2 clean
manuscript, rendered as Word insertions and deletions.

Run after the markdown sources are finalized:
  python3 make_docx.py
"""
from __future__ import annotations
import re
import sys
import difflib
import subprocess
from pathlib import Path

NB = Path(__file__).resolve().parents[4] / "notebooks" / "panelfm"
DOCS = NB / "revision2" / "docs"
R1_CLEAN = NB / "revision1" / "docs" / "manuscript_blinded_clean.md"
REF = Path.home() / ".claude" / "templates" / "sanjay_paper_reference.docx"


def strip_bold(text: str) -> str:
    # Remove the revision bold markup (**...**) but keep other markdown intact.
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text, flags=re.DOTALL)


def pandoc(md_path: Path, docx_path: Path):
    cmd = ["pandoc", str(md_path), "--from=markdown", "--to=docx", "-o", str(docx_path)]
    if REF.exists():
        cmd.append(f"--reference-doc={REF}")
    subprocess.run(cmd, check=True)
    print(f"  wrote {docx_path.name}")


def build_tracked_changes(r1_md_path: Path, r2_clean_md_path: Path, out_path: Path):
    """Paragraph-level R1->R2 redline that preserves Word formatting.

    Both versions are rendered to docx with pandoc (so headings, bold, and the
    reference-template styles are real Word formatting, not markdown source), then
    aligned paragraph-by-paragraph; inserted and deleted paragraphs' runs are
    wrapped in native Word tracked-changes (w:ins / w:del) so Word shows a redline.
    """
    import copy
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    import tempfile
    tmp = Path(tempfile.mkdtemp())
    r1_docx, r2_docx = tmp / "r1.docx", tmp / "r2.docx"
    pandoc(r1_md_path, r1_docx)
    pandoc(r2_clean_md_path, r2_docx)

    d1, d2 = Document(str(r1_docx)), Document(str(r2_docx))
    p1, p2 = d1.paragraphs, d2.paragraphs
    norm = lambda s: re.sub(r"\s+", " ", s).strip()
    t1 = [norm(p.text) for p in p1]
    t2 = [norm(p.text) for p in p2]

    AUTHOR, DATE = "Revision 2", "2026-06-17T00:00:00Z"
    rid = [0]

    def _wrap_ins(p_elem):
        for r in p_elem.findall(qn("w:r")):
            rid[0] += 1
            ins = OxmlElement("w:ins")
            ins.set(qn("w:id"), str(rid[0])); ins.set(qn("w:author"), AUTHOR); ins.set(qn("w:date"), DATE)
            p_elem.replace(r, ins); ins.append(r)

    def _wrap_del(p_elem):
        for r in p_elem.findall(qn("w:r")):
            rid[0] += 1
            d = OxmlElement("w:del")
            d.set(qn("w:id"), str(rid[0])); d.set(qn("w:author"), AUTHOR); d.set(qn("w:date"), DATE)
            for t in r.findall(qn("w:t")):  # deleted text uses w:delText
                dt = OxmlElement("w:delText"); dt.set(qn("xml:space"), "preserve"); dt.text = t.text
                r.replace(t, dt)
            p_elem.replace(r, d); d.append(r)

    body = d2.element.body
    sectPr = body.find(qn("w:sectPr"))
    sm = difflib.SequenceMatcher(a=t1, b=t2, autojunk=False)
    ins_idx = set()
    del_before = {}  # R2 paragraph index -> list of R1 paragraph indices to insert (as deletions) before it
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag in ("insert", "replace"):
            ins_idx.update(range(j1, j2))
        if tag in ("delete", "replace"):
            del_before.setdefault(j1, []).extend(range(i1, i2))

    for j in ins_idx:
        _wrap_ins(p2[j]._p)

    for j, i_list in sorted(del_before.items()):
        anchor = p2[j]._p if j < len(p2) else sectPr
        for i in i_list:
            del_p = copy.deepcopy(p1[i]._p)
            _wrap_del(del_p)
            if anchor is not None:
                anchor.addprevious(del_p)
            else:
                body.append(del_p)

    d2.save(str(out_path))
    print(f"  wrote {out_path.name} ({rid[0]} tracked revisions, formatting preserved)")


def main():
    bold_md = DOCS / "manuscript_bold.md"
    if not bold_md.exists():
        print(f"ERROR: {bold_md} not found", file=sys.stderr); sys.exit(1)
    bold_text = bold_md.read_text()
    clean_text = strip_bold(bold_text)
    clean_md = DOCS / "manuscript_clean.md"
    clean_md.write_text(clean_text)

    print("Building manuscript docx...")
    pandoc(bold_md, DOCS / "manuscript_bold.docx")
    pandoc(clean_md, DOCS / "manuscript_clean.docx")

    appx = DOCS / "supplementary_appendix.md"
    if appx.exists():
        pandoc(appx, DOCS / "supplementary_appendix.docx")
    resp = DOCS / "response_to_reviewers.md"
    if resp.exists():
        pandoc(resp, DOCS / "response_to_reviewers.docx")
    title = DOCS / "title_page.md"
    if title.exists():
        pandoc(title, DOCS / "title_page.docx")

    print("Building tracked-changes redline...")
    if R1_CLEAN.exists():
        build_tracked_changes(R1_CLEAN, clean_md, DOCS / "manuscript_tracked_changes.docx")
    else:
        print(f"  WARNING: {R1_CLEAN} not found; skipping redline")

    print("Done.")


if __name__ == "__main__":
    main()
