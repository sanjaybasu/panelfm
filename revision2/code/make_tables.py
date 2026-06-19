#!/usr/bin/env python3
"""
Build editable Word tables for the second revision:
  table1_demographics.docx  cohort and disjoint member-set characteristics
  table2_performance.docx    model performance with bootstrap 95% CIs

Run after run_member_disjoint.py has produced the results JSONs.
"""
import os
import sys
import json
from pathlib import Path

import numpy as np
import pandas as pd

THIS_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(THIS_DIR))
import run_member_disjoint as R
import docx
from docx.shared import Pt

NB_TABLES = Path(__file__).resolve().parents[4] / "notebooks" / "panelfm" / "revision2" / "tables"
NB_TABLES.mkdir(parents=True, exist_ok=True)
RES = R.OUT_DIR


def median_iqr(x):
    x = np.asarray(x, dtype=float)
    return f"{np.median(x):.0f} ({np.percentile(x,25):.0f}-{np.percentile(x,75):.0f})"


def set_stats(outcomes, members, label):
    sub = outcomes[outcomes["person_id"].isin(members)]
    pm = sub.groupby("person_id").size()
    cost = sub["total_paid"].astype(float)
    return {
        "Set": label,
        "Members": f"{sub['person_id'].nunique():,}",
        "Patient-months": f"{sub.shape[0]:,}",
        "Months/member, median (IQR)": median_iqr(pm.values),
        "Monthly cost $, median (IQR)": f"{cost.median():.0f} ({cost.quantile(.25):.0f}-{cost.quantile(.75):.0f})",
        "Monthly cost $, mean (SD)": f"{cost.mean():.0f} ({cost.std():.0f})",
        "$0-cost months, %": f"{100*(cost==0).mean():.1f}",
    }


def build_table1():
    outcomes, attributes, eligibility = R.load_cohort()
    data = R.prepare(outcomes, attributes, eligibility)
    sp = data["splits"]
    oc = data["outcomes"]
    rows = [
        set_stats(oc, set(oc["person_id"].unique()), "Full cohort"),
        set_stats(oc, sp["train"], "Training members (disjoint)"),
        set_stats(oc, sp["val"], "Validation members (disjoint)"),
        set_stats(oc, sp["test"], "Test members (disjoint)"),
    ]
    df = pd.DataFrame(rows)
    df.to_csv(NB_TABLES / "table1_demographics.csv", index=False)

    doc = docx.Document()
    doc.add_paragraph("Table 1. Cohort and member-disjoint set characteristics.")
    t = doc.add_table(rows=1, cols=len(df.columns)); t.style = "Light Grid Accent 1"
    for j, c in enumerate(df.columns):
        t.rows[0].cells[j].text = c
    for _, r in df.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(df.columns):
            cells[j].text = str(r[c])
    doc.add_paragraph(
        "Members were partitioned into disjoint training, validation, and test sets (stratified by "
        "lookback-period cost); no member appears in more than one set. The full cohort row includes all "
        "enrolled members; the disjoint-set rows include only the 57,367 members meeting prospective-forecast "
        "inclusion criteria. The three disjoint sets are distributionally similar by construction.")
    doc.save(str(NB_TABLES / "table1_demographics.docx"))
    print("wrote table1_demographics.docx")


def fmt_ci(ci):
    return f"{ci['mae']:.0f} ({ci['mae_ci'][0]:.0f}-{ci['mae_ci'][1]:.0f})", \
           f"{ci['r2_cal']:.2f} ({ci['r2_cal_ci'][0]:.2f} to {ci['r2_cal_ci'][1]:.2f})"


def build_table2():
    m = json.loads((RES / "all_metrics_disjoint.json").read_text())
    ci = json.loads((RES / "bootstrap_ci_disjoint.json").read_text())

    prospective = [
        ("Demographics-only GLM", "demographic_glm"),
        ("XGBoost", "xgboost"), ("Random forest", "random_forest"),
        ("LightGBM", "lightgbm"), ("Two-part hurdle", "two_part"), ("Stacking ensemble", "stacking"),
        ("Naive trailing mean", "naive_mean3"), ("Naive last-value", "naive_last"),
        ("Chronos (zero-shot)", "chronos_zeroshot"), ("PanelFM-XReg", "panelfm_xreg"),
        ("PanelFM-Adapter", "panelfm_adapter"),
        ("Hybrid (proportional)", "hybrid_chronos_zeroshot"), ("Hybrid (gated)", "hybrid_gated"),
    ]
    concurrent = [("Concurrent XGBoost", "concurrent_xgboost"), ("Concurrent random forest", "concurrent_random_forest"),
                  ("Concurrent LightGBM", "concurrent_lightgbm"), ("Concurrent two-part", "concurrent_two_part")]

    cols = ["Model", "MAE $/mbr-mo (95% CI)", "Calibrated R² (95% CI)", "Predictive ratio", "c-statistic", "PPV top 10%", "Lift"]

    def row_for(label, key):
        d = m[key]
        if key in ci:
            mae_s, r2_s = fmt_ci(ci[key])
        else:
            mae_s = f"{d['mae']:.0f}"
            r2_s = f"{d['r_squared_calibrated']:.2f}"
        return [label, mae_s, r2_s, f"{d['predictive_ratio']:.2f}",
                f"{(d.get('auroc') or float('nan')):.2f}", f"{(d.get('ppv_at_top10pct') or float('nan')):.2f}",
                f"{(d.get('lift') or float('nan')):.1f}"]

    doc = docx.Document()
    doc.add_paragraph("Table 2. Prospective and concurrent model performance on held-out test members (n = 8,604).")
    t = doc.add_table(rows=1, cols=len(cols)); t.style = "Light Grid Accent 1"
    for j, c in enumerate(cols):
        t.rows[0].cells[j].text = c

    def section(title):
        cells = t.add_row().cells
        cells[0].merge(cells[len(cols) - 1])
        cells[0].text = title
        for r in cells[0].paragraphs:
            for run in r.runs:
                run.bold = True

    section("Prospective models")
    for label, key in prospective:
        if key in m:
            cells = t.add_row().cells
            for j, v in enumerate(row_for(label, key)):
                cells[j].text = v
    section("Concurrent models (explanatory ceiling, not forecasts)")
    for label, key in concurrent:
        if key in m:
            cells = t.add_row().cells
            for j, v in enumerate(row_for(label, key)):
                cells[j].text = v

    doc.add_paragraph(
        "MAE is per member-month; calibrated R² and the predictive ratio are computed on the "
        "patient three-month total. 95% confidence intervals are from 2,000-iteration member-level bootstrap "
        "resampling. The gated hybrid attains lower MAE than the mean-targeting cross-sectional models while "
        "retaining positive between-member discrimination. Concurrent models use same-period features and represent an explanatory ceiling, not "
        "prospective forecasting performance.")
    doc.save(str(NB_TABLES / "table2_performance.docx"))
    print("wrote table2_performance.docx")


if __name__ == "__main__":
    build_table1()
    build_table2()
    print("done")
